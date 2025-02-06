import os
import docx
import re

def extract_text_from_docx(docx_path, formats):
    doc = docx.Document(docx_path)
    text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
            formats.add("text")
    
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(row_text))
        text.append("\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]\n")
        formats.add("table")
    
    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            text.append("[รูปภาพ]")
            formats.add("image")
    
    return '\n'.join(text)

def process_docx_files(input_folder, output_folder, overwrite=False):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    if not overwrite:
        processed_files = os.listdir(output_folder)
    else:
        processed_files = []
    
    formats_found = set()
    
    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):
            input_path = os.path.join(input_folder, filename)
            
            match = re.search(r'(?i)kb\s*\d+', filename)
            if match:
                new_filename = f"{match.group().replace(' ', '').upper()}.txt"
            else:
                new_filename = f"processed_{filename}.txt"
            
            output_path = os.path.join(output_folder, new_filename)
            
            text = extract_text_from_docx(input_path, formats_found)
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)
            
            print(f"Processed: {filename} -> {output_path}")
    
    print("\nSummary of content types found in KB folder:")
    for format_type in formats_found:
        print(f"- {format_type}")
    fileNames = os.listdir(output_folder)
    # save the file names to a text file
    with open('processed_files.txt', 'w') as f:
        text = ",".join(fileNames).replace(".txt", "")
        f.write(text)

# Define input and output directories
input_folder = "KB"
output_folder = "processed_KB"

# Process all docx files
process_docx_files(input_folder, output_folder, overwrite=True)
